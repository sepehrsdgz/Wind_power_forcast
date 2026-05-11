from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from path import win_to_wsl as ws


# ==========================================
# 1. HELPER FUNCTIONS (Formulas & Formatting)
# ==========================================

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_equation_xml(paragraph, xml_str):
    """
    Injects raw OXML (Office Math Markup Language) into a paragraph.
    This creates a real, editable Word equation.
    """
    # Get the paragraph's underlying XML element
    p_element = paragraph._p

    # Create a wrapper for the OMML
    # We parse the XML string into an element and append it
    try:
        # We need to wrap the user's math XML in the standard namespaces
        # However, OxmlElement.fromstring is not directly exposed in simple imports sometimes.
        # We will assume xml_str is a fully formed <m:oMath> or <m:oMathPara> snippet.
        from docx.oxml import parse_xml
        math_element = parse_xml(xml_str)
        p_element.append(math_element)
    except Exception as e:
        print(f"Error adding equation: {e}")
        # Fallback to text if XML fails
        paragraph.add_run(" [Formula Error] ")


def add_bold_paragraph(doc, bold_text, normal_text):
    """Helper to add 'Bold text: Normal text' without using ** syntax"""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(normal_text)


# --- RAW XML FOR YOUR FORMULAS ---
# These are pre-compiled OMML (Office Math) strings for the specific LaTeX codes you provided.

# 1. Power Law: v = v_0 * (h/h_0)^alpha
xml_power_law = (
    r'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
    r'<m:oMath><m:r><m:t>v</m:t></m:r><m:r><m:t>=</m:t></m:r>'
    r'<m:sSub><m:e><m:r><m:t>v</m:t></m:r></m:e><m:sub><m:r><m:t>0</m:t></m:r></m:sub></m:sSub>'
    r'<m:sSup><m:e><m:d><m:dPr><m:ctrlPr/></m:dPr><m:e><m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
    r'<m:num><m:r><m:t>h</m:t></m:r></m:num><m:den><m:sSub><m:e><m:r><m:t>h</m:t></m:r></m:e>'
    r'<m:sub><m:r><m:t>0</m:t></m:r></m:sub></m:sSub></m:den></m:f></m:e></m:d></m:e>'
    r'<m:sup><m:r><m:t>α</m:t></m:r></m:sup></m:sSup></m:oMath></m:oMathPara>'
)

# 2. Cubic Power: P_T(v) approx a*v^3 - b*P_R
xml_cubic_power = (
    r'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
    r'<m:oMath><m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>T</m:t></m:r></m:sub></m:sSub>'
    r'<m:r><m:t>(v)≈a∙</m:t></m:r><m:sSup><m:e><m:r><m:t>v</m:t></m:r></m:e><m:sup><m:r><m:t>3</m:t></m:r></m:sup></m:sSup>'
    r'<m:r><m:t>−</m:t></m:r><m:r><m:t>b∙</m:t></m:r>'
    r'<m:sSub><m:e><m:r><m:t>P</m:t></m:r></m:e><m:sub><m:r><m:t>R</m:t></m:r></m:sub></m:sSub>'
    r'</m:oMath></m:oMathPara>'
)

# 3. RMSE
xml_rmse = (
    r'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath>'
    r'<m:r><m:t>RMSE</m:t></m:r><m:r><m:t>=</m:t></m:r>'
    r'<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr><m:deg/><m:e>'
    r'<m:f><m:fPr><m:type m:val="bar"/></m:fPr><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>N</m:t></m:r></m:den></m:f>'
    r'<m:nary><m:naryPr><m:chr m:val="∑"/><m:limitLoc m:val="undOvr"/><m:subHide m:val="1"/><m:supHide m:val="1"/></m:naryPr>'
    r'<m:sub/><m:sup/><m:e><m:sSup><m:e><m:d><m:dPr><m:ctrlPr/></m:dPr><m:e>'
    r'<m:sSub><m:e><m:r><m:t>y</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'<m:r><m:t>−</m:t></m:r>'
    r'<m:sSub><m:e><m:acc><m:accPr><m:chr m:val="̂"/></m:accPr><m:e><m:r><m:t>y</m:t></m:r></m:e></m:acc></m:e>'
    r'<m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'</m:e></m:d></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:e></m:nary>'
    r'</m:e></m:rad></m:oMath></m:oMathPara>'
)

# 4. MAE
xml_mae = (
    r'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath>'
    r'<m:r><m:t>MAE</m:t></m:r><m:r><m:t>=</m:t></m:r>'
    r'<m:f><m:fPr><m:type m:val="bar"/></m:fPr><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>N</m:t></m:r></m:den></m:f>'
    r'<m:nary><m:naryPr><m:chr m:val="∑"/><m:limitLoc m:val="undOvr"/><m:subHide m:val="1"/><m:supHide m:val="1"/></m:naryPr>'
    r'<m:sub/><m:sup/><m:e><m:d><m:dPr><m:begChr m:val="|"/><m:endChr m:val="|"/></m:dPr><m:e>'
    r'<m:sSub><m:e><m:r><m:t>y</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'<m:r><m:t>−</m:t></m:r>'
    r'<m:sSub><m:e><m:acc><m:accPr><m:chr m:val="̂"/></m:accPr><m:e><m:r><m:t>y</m:t></m:r></m:e></m:acc></m:e>'
    r'<m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'</m:e></m:d></m:e></m:nary></m:oMath></m:oMathPara>'
)

# 5. R^2
xml_r2 = (
    r'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath>'
    r'<m:sSup><m:e><m:r><m:t>R</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>'
    r'<m:r><m:t>=1−</m:t></m:r>'
    r'<m:f><m:fPr><m:type m:val="bar"/></m:fPr><m:num>'
    r'<m:nary><m:naryPr><m:chr m:val="∑"/><m:limitLoc m:val="undOvr"/><m:subHide m:val="1"/><m:supHide m:val="1"/></m:naryPr>'
    r'<m:sub/><m:sup/><m:e><m:sSup><m:e><m:d><m:dPr><m:ctrlPr/></m:dPr><m:e>'
    r'<m:sSub><m:e><m:r><m:t>y</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'<m:r><m:t>−</m:t></m:r>'
    r'<m:sSub><m:e><m:acc><m:accPr><m:chr m:val="̂"/></m:accPr><m:e><m:r><m:t>y</m:t></m:r></m:e></m:acc></m:e>'
    r'<m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'</m:e></m:d></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:e></m:nary>'
    r'</m:num><m:den>'
    r'<m:nary><m:naryPr><m:chr m:val="∑"/><m:limitLoc m:val="undOvr"/><m:subHide m:val="1"/><m:supHide m:val="1"/></m:naryPr>'
    r'<m:sub/><m:sup/><m:e><m:sSup><m:e><m:d><m:dPr><m:ctrlPr/></m:dPr><m:e>'
    r'<m:sSub><m:e><m:r><m:t>y</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>'
    r'<m:r><m:t>−</m:t></m:r>'
    r'<m:acc><m:accPr><m:chr m:val="̄"/></m:accPr><m:e><m:r><m:t>y</m:t></m:r></m:e></m:acc>'
    r'</m:e></m:d></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:e></m:nary>'
    r'</m:den></m:f></m:oMath></m:oMathPara>'
)

# ==========================================
# 2. DOCUMENT SETUP
# ==========================================
# PATHS
formula_image_path1 = ws(r"Figure1_PowerCurve.png")
formula_image_path2 = ws(r"Figure2_WindCharacteristics.png")
formula_image_path3 = ws(r"Figure3_FrequencyDistribution.png")

case1_image_path = ws(
    r"\\wsl.localhost\Ubuntu-20.04\home\sepehr\ml\Wind_power_forcast\Final_Plots\Fig6a_Scatter_LightGBM_Tuned.png")
case2_image_path = ws(
    r"\\wsl.localhost\Ubuntu-20.04\home\sepehr\ml\Wind_power_forcast\Case2_Results\Case2_Scatter_LightGBM_Tuned.png")
case3_image_path = ws(
    r"\\wsl.localhost\Ubuntu-20.04\home\sepehr\ml\Wind_power_forcast\Final_Report_Outputs\Scatter_Transfer_LightGBM_WS_STD.png")

doc = Document()

# ==========================================
# 3. METHODOLOGY CONTENT
# ==========================================
doc.add_heading('3. Methodology', level=1)

p = doc.add_paragraph(
    "This study adopts and extends the methodological framework proposed by Demolli et al. [1] to forecast long-term wind power generation. "
    "While the original study focused on locations in Turkey, this research applies the methodology to a new geographical context (Chicago and Detroit, USA) "
    "to test the generalizability of the proposed machine learning models. Furthermore, this study introduces a modern algorithmic extension "
    "(LightGBM) and a robust Bayesian optimization strategy (Optuna) to improve upon the original trial-and-error approach."
)

# --- 3.1 Data Acquisition ---
doc.add_heading('3.1 Data Acquisition and Site Selection', level=2)

doc.add_paragraph(
    "Hourly meteorological data was obtained from the open-access \"Historical Hourly Weather Data 2012-2017\" dataset available on Kaggle [3]. "
    "The city of Chicago was selected as the primary \"Base Location\" for training and initial validation due to its consistent wind characteristics. "
    "Additionally, data for a second city, Detroit, was acquired to serve as an independent test set for evaluating the geographical transferability of the models."
)
doc.add_paragraph(
    "The dataset spans exactly five years, from October 1, 2012, to October 1, 2017. Following the protocol of the original study, "
    "the first four years (80%) were utilized for training and validation, while the final year (20%) was reserved strictly for testing."
)

# Insert Figures 2 & 3
try:
    doc.add_picture(formula_image_path2, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 2: Wind Speed Characteristics", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

try:
    doc.add_picture(formula_image_path3, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 3: Weibull Frequency Distribution", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# --- 3.2 Theoretical Power ---
doc.add_heading('3.2 Theoretical Power Generation (Target Synthesis)', level=2)
doc.add_paragraph(
    "A critical challenge in wind energy forecasting is the lack of public datasets containing both meteorological inputs and actual turbine power output. "
    "To address this, the target variable (Power Output) was synthesized mathematically based on the physics of a specific wind turbine."
)

doc.add_heading('3.2.1 Vertical Extrapolation', level=3)
doc.add_paragraph(
    "The raw meteorological data provided wind speeds at a standard height of 10 meters, which corresponds to the measurement standards used by "
    "services such as the OpenWeatherMap API [4]. However, commercial wind turbines operate at significantly higher altitudes. "
    "Therefore, the wind speed was extrapolated to a hub height of 50 meters using the Power Law equation:"
)

# --- EQUATION 1 (Power Law) ---
p = doc.add_paragraph()
add_equation_xml(p, xml_power_law)

doc.add_paragraph(
    "Where v is the wind speed at hub height (50m), v0 is the measured speed at 10m, and alpha is the Hellman exponent (set to 0.14 for open terrain)."
)

doc.add_heading('3.2.2 The Cubic Power Curve Model', level=3)
doc.add_paragraph(
    "In the reference study by Demolli et al., the power generation data was derived using specific manufacturer power/speed tables. "
    "In contrast, this study utilizes a theoretical approach to ensure replicability. The relationship between wind speed and power generation "
    "was modeled using the \"Four-Region\" logic described in the course lecture notes by Bayindir [2]."
)
doc.add_paragraph(
    "Specifically, for the operational range (Region 2), the power output is calculated using the cubic formula visualized in Figure 1:"
)

# --- EQUATION 2 (Power Curve) ---
p = doc.add_paragraph()
add_equation_xml(p, xml_cubic_power)

try:
    doc.add_picture(formula_image_path1, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 1: Theoretical Power Curve", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph(
    "This formula ensures that the power output follows the physical cubic law of kinetic energy. "
    "The turbine parameters were set to: Cut-in (3 m/s), Rated (15 m/s), and Cut-out (25 m/s)."
)

# --- 3.3 Models ---
doc.add_heading('3.3 Machine Learning Models and Hyperparameters', level=2)
doc.add_paragraph(
    "Six regression algorithms were implemented to model the relationship between wind statistics and power generation. "
    "To strictly validate the findings of the original study, the hyperparameters for the five benchmark models were kept identical "
    "to those reported by Demolli et al. [1]. No additional parameter optimization was performed for these specific models."
)

# Fixed bold syntax
doc.add_paragraph("The fixed configuration for the benchmark models is as follows:")
doc.add_paragraph("LASSO Regression: alpha=0.1", style='List Bullet')
doc.add_paragraph("k-Nearest Neighbors (kNN): k=4, Metric=Minkowski", style='List Bullet')
doc.add_paragraph("Random Forest (RF): n_estimators=10", style='List Bullet')
doc.add_paragraph("XGBoost: n_estimators=500, learning_rate=0.1", style='List Bullet')
doc.add_paragraph("Support Vector Regression (SVR): C=3000, gamma=0.1, epsilon=0.1, kernel=RBF", style='List Bullet')

add_bold_paragraph(doc, "Algorithmic Extension (LightGBM): ",
                   "The sixth model, LightGBM, was not present in the original study. "
                   "Its hyperparameters were optimized using Optuna, a Bayesian optimization framework."
                   )

# --- 3.4 Experimental Design ---
doc.add_heading('3.4 Experimental Design and Case Studies', level=2)
doc.add_paragraph(
    "To comprehensively evaluate the models, the methodology was structured into three distinct experimental cases.")

# Case 1
doc.add_heading('Case 1: Baseline Forecasting (Chicago)', level=3)
add_bold_paragraph(doc, "Objective: ",
                   "To establish the baseline performance of the algorithms using the complete feature set.")
add_bold_paragraph(doc, "Inputs: ", "Daily Mean Wind Speed + Daily Standard Deviation.")
add_bold_paragraph(doc, "Methodology: ",
                   "Training and testing were performed on the Chicago dataset. This duplicates 'Case 1' from the reference paper.")

# Case 2
doc.add_heading('Case 2: Input Sensitivity Analysis', level=3)
add_bold_paragraph(doc, "Objective: ", "To quantify the impact of wind volatility (turbulence) on prediction accuracy.")
add_bold_paragraph(doc, "Inputs: ", "Daily Mean Wind Speed only.")
add_bold_paragraph(doc, "Methodology: ",
                   "The 'Standard Deviation' feature was removed. This mirrors 'Case 2' of the original study.")

# Case 3
doc.add_heading('Case 3: Geographical Transferability', level=3)
add_bold_paragraph(doc, "Objective: ",
                   "To assess the generalization capability of the models when applied to a new location.")
add_bold_paragraph(doc, "Inputs: ", "Daily Mean Wind Speed + Daily Standard Deviation.")
add_bold_paragraph(doc, "Methodology: ",
                   "Models trained on Chicago were frozen and used to predict power generation in Detroit.")

# --- 3.6 Metrics ---
doc.add_heading('3.6 Performance Metrics', level=2)
doc.add_paragraph(
    "To quantitatively evaluate the forecasting accuracy, three standard statistical metrics were employed:")

# RMSE
doc.add_paragraph("Root Mean Squared Error (RMSE):", style='List Bullet')
p = doc.add_paragraph()
add_equation_xml(p, xml_rmse)

# MAE
doc.add_paragraph("Mean Absolute Error (MAE):", style='List Bullet')
p = doc.add_paragraph()
add_equation_xml(p, xml_mae)

# R2
doc.add_paragraph("Coefficient of Determination (R²):", style='List Bullet')
p = doc.add_paragraph()
add_equation_xml(p, xml_r2)

# ==========================================
# 4. RESULTS
# ==========================================
doc.add_heading('4. Results', level=1)
doc.add_paragraph(
    "This section presents the performance of the proposed algorithms across the three experimental cases.")

doc.add_heading('4.1 Case 1 Results (Chicago: Mean + Std Dev)', level=2)
doc.add_paragraph(
    "Table 1 summarizes the forecasting accuracy for Case 1. XGBoost and SVR achieved the highest performance.")

try:
    doc.add_picture(case1_image_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 4: Case 1 Time Series Forecast (SVR)",
                      style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_heading('4.2 Case 2 Results (Chicago: Mean Only)', level=2)
doc.add_paragraph("Table 2 presents the results when the standard deviation feature is removed.")

try:
    doc.add_picture(case2_image_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 5: Case 2 Scatter Plot (Wind Speed vs Power)",
                      style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_heading('4.3 Case 3 Results (Transferability to Detroit)', level=2)
doc.add_paragraph(
    "Table 3 compares the transferability of the models when applied to the Detroit dataset (Blind Test).")

try:
    doc.add_picture(case3_image_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 6: Case 3 Transferability Forecast (Detroit)",
                      style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# ==========================================
# 5. REFERENCES
# ==========================================
doc.add_heading('5. References', level=1)
refs = [
    "[1] H. Demolli, A. S. Dokuz, A. Ecemis, and M. Gokcek, \"Wind power forecasting based on daily wind speed data using machine learning algorithms,\" Energy Convers. Manag., vol. 198, p. 111823, Oct. 2019.",
    "[2] R. Bayindir, \"Wind Energy Systems Lecture Notes,\" Department of Electrical Engineering, Gazi University. [Online].",
    "[3] SelfishGene, \"Historical Hourly Weather Data 2012-2017,\" Kaggle. [Online]. Available: https://www.kaggle.com/datasets/selfishgene/historical-hourly-weather-data.",
    "[4] OpenWeatherMap, \"Weather API - Current weather and forecast,\" OpenWeatherMap. [Online]. Available: https://openweathermap.org/api/weather-map-2."
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after = Pt(6)

# Save
file_name = "Methodology_Report_Formulas_Final.docx"
doc.save(file_name)
print(f"Document saved as {file_name}")