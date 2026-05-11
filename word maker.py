from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('3. Methodology', level=1)

# Intro
p = doc.add_paragraph(
    "This study adopts and extends the methodological framework proposed by Demolli et al. [1] to forecast long-term wind power generation. "
    "While the original study focused on locations in Turkey, this research applies the methodology to a new geographical context (Chicago, USA) "
    "to test the generalizability of the proposed machine learning models. Furthermore, this study introduces a modern algorithmic extension "
    "(LightGBM) and a robust Bayesian optimization strategy to improve upon the original trial-and-error approach."
)

# 3.1 Data Acquisition
doc.add_heading('3.1 Data Acquisition and Site Selection', level=2)
doc.add_paragraph(
    "Hourly meteorological data was obtained from the open-access \"Historical Hourly Weather Data 2012-2017\" dataset. "
    "The city of Chicago was selected as the primary case study due to its consistent wind characteristics and the completeness of the time-series records."
)
doc.add_paragraph(
    "The dataset spans exactly five years, from October 1, 2012, to October 1, 2017. Following the protocol of the original study, "
    "the first four years (80%) were utilized for training and validation, while the final year (20%) was reserved strictly for testing."
)

# 3.2 Theoretical Power
doc.add_heading('3.2 Theoretical Power Generation (Target Synthesis)', level=2)
doc.add_paragraph(
    "A critical challenge in wind energy forecasting is the lack of public datasets containing both meteorological inputs and actual turbine power output. "
    "To address this, the target variable (Power Output) was synthesized mathematically based on the physics of a specific wind turbine."
)

doc.add_heading('3.2.1 Vertical Extrapolation', level=3)
doc.add_paragraph(
    "The raw meteorological data provided wind speeds at a standard height of 10 meters. However, commercial wind turbines operate at significantly higher altitudes. "
    "Therefore, the wind speed was extrapolated to a hub height of 50 meters using the Power Law equation described in Demolli et al. [1]:"
)
# Equation 1
doc.add_paragraph("v = v0 * (h / h0)^alpha", style='Quote')
doc.add_paragraph(
    "Where v is the wind speed at hub height (50m), v0 is the measured speed at 10m, and alpha is the Hellman exponent (set to 0.14 for open terrain)."
)

doc.add_heading('3.2.2 The Cubic Power Curve Model', level=3)
doc.add_paragraph(
    "To generate the \"Ground Truth\" power values, a theoretical 1 MW wind turbine was simulated. The operational characteristics were defined based on "
    "Table 1 of the original study (Cut-in: 3 m/s, Rated: 15 m/s, Cut-out: 25 m/s)."
)
doc.add_paragraph(
    "The relationship between wind speed and power generation was modeled using the \"Four-Region\" logic described in the course lecture notes by Bayindir [2]. "
    "Specifically, for the ramp-up phase (Region 4), Equation 9.82 was utilized to model the cubic increase in power:"
)
# Equation 2
doc.add_paragraph("PT(v) ≈ a * v^3 - b * PR", style='Quote')
doc.add_paragraph(
    "This formula ensures that the power output follows the physical v^3 law of kinetic energy, providing a realistic target variable for the regression algorithms."
)

# 3.3 Data Preprocessing
doc.add_heading('3.3 Data Preprocessing and Feature Engineering', level=2)
doc.add_paragraph("The raw hourly data was aggregated into daily samples to match the temporal resolution of the original study. For each day, two input features were engineered:")
doc.add_paragraph("1. Daily Mean Wind Speed: Representing the average energy potential.", style='List Number')
doc.add_paragraph("2. Daily Standard Deviation: Representing the turbulence or volatility of the wind.", style='List Number')
doc.add_paragraph("The target variable was the Daily Total Power, calculated by summing the synthesized hourly power output for each 24-hour period.")

# 3.4 ML Algorithms
doc.add_heading('3.4 Machine Learning Algorithms', level=2)
doc.add_paragraph("Six regression algorithms were implemented to model the relationship between wind statistics and power generation.")

doc.add_heading('3.4.1 Replication Models', level=3)
doc.add_paragraph(
    "To verify the findings of Demolli et al., five algorithms were trained using the exact hyperparameters specified in the original paper. "
    "This ensures that any deviation in performance is due to data characteristics rather than parameter tuning:"
)
# List of models
models = [
    "LASSO Regression: (alpha=0.1) Used as a linear baseline.",
    "k-Nearest Neighbors (kNN): (k=4, Minkowski distance).",
    "Random Forest (RF): (n=10 trees).",
    "XGBoost: (n=500, learning rate=0.1).",
    "Support Vector Regression (SVR): (C=3000, gamma=0.1, RBF kernel)."
]
for model in models:
    doc.add_paragraph(model, style='List Bullet')

doc.add_heading('3.4.2 Algorithmic Extension: LightGBM', level=3)
doc.add_paragraph(
    "This study extends the original methodology by introducing LightGBM (Light Gradient Boosting Machine). "
    "LightGBM utilizes a leaf-wise tree growth strategy and Gradient-based One-Side Sampling (GOSS), which theoretically offers faster convergence "
    "and higher accuracy on non-linear datasets compared to the level-wise growth of XGBoost."
)

# 3.5 Optimization
doc.add_heading('3.5 Hyperparameter Optimization Strategy', level=2)
doc.add_paragraph(
    "The original study relied on a manual trial-and-error approach for parameter selection. To establish a more rigorous baseline for the new LightGBM model, "
    "this study implemented Bayesian Optimization using the Optuna framework."
)
doc.add_paragraph(
    "To prevent \"data leakage\"—where a model inadvertently learns from the test set—a Time-Series Cross-Validation strategy was employed. "
    "The training data (Years 1-4) was split into three chronological folds. The optimization algorithm minimized the Root Mean Squared Error (RMSE) "
    "across these folds, ensuring that the selected parameters were robust against temporal variations before being applied to the final test year."
)

# 3.6 Metrics
doc.add_heading('3.6 Performance Metrics', level=2)
doc.add_paragraph(
    "Since the objective is to predict a continuous variable, classification metrics (Precision/Recall) were deemed inappropriate. "
    "The models were evaluated using the following regression metrics:"
)
metrics = [
    "Coefficient of Determination (R^2): To measure the proportion of variance explained by the model.",
    "Root Mean Squared Error (RMSE): To quantify the magnitude of prediction errors, penalizing larger deviations.",
    "Mean Absolute Error (MAE): To represent the average error in kilowatts (kW)."
]
for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

# Save
file_name = "Methodology_Section.docx"
doc.save(file_name)
print(f"Document saved as {file_name}")